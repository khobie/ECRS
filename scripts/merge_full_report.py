"""Append Chapters 4 and 5 to the original report docx."""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

# Import builder from sibling script
import sys
sys.path.insert(0, str(Path(__file__).parent))
from build_chapters_docx import setup_styles, build_chapter_four_five

ORIGINAL = Path(r"c:\Users\pc\Downloads\Design_and_Implementation_of_a_Crime_Reporting_Web_Application_chpt_1_2_3 (2).docx")
OUT = Path(__file__).resolve().parents[1] / "docs" / "FULL_REPORT_CHAPTERS_1_to_5.docx"


def main():
    if not ORIGINAL.exists():
        raise FileNotFoundError(f"Original not found: {ORIGINAL}")

    doc = Document(str(ORIGINAL))
    setup_styles(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    build_chapter_four_five(doc)
    doc.save(OUT)
    print(f"Created full report: {OUT}")


if __name__ == "__main__":
    main()
