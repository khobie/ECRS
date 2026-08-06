"""Embed all screenshots (frontend, backend, database) into Word report."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
SHOTS = ROOT / "docs" / "screenshots"
SRC = ROOT / "docs" / "FULL_REPORT_CHAPTERS_1_to_5.docx"
OUT = ROOT / "docs" / "FULL_REPORT_WITH_SCREENSHOTS.docx"
DL = Path(r"c:\Users\pc\Downloads\FULL_REPORT_WITH_SCREENSHOTS_v2.docx")

SECTIONS = [
    ("FRONTEND SCREENSHOTS", [
        ("fig-4-1-landing.png", "Figure 4.1: ECRS Landing Page (Citizen Portal)"),
        ("fig-4-2-report-step1.png", "Figure 4.2: Crime Report Submission – Incident Step"),
        ("fig-4-4-track-report.png", "Figure 4.3: Track Report Page Showing Case Status"),
        ("fig-4-5-officer-login.png", "Figure 4.4: Officer Login Page (Separate Portal)"),
        ("fig-4-6-dashboard.png", "Figure 4.5: Officer Dashboard with Live Statistics"),
        ("fig-4-7-investigation.png", "Figure 4.6: Case Investigation View"),
        ("fig-4-8-analytics.png", "Figure 4.7: Crime Analytics Page"),
    ]),
    ("BACKEND (LARAVEL API) SCREENSHOTS", [
        ("fig-backend-health.png", "Figure 4.8: Backend API – Health Check (GET /api/health)"),
        ("fig-backend-zones.png", "Figure 4.9: Backend API – Zones Endpoint (GET /api/zones)"),
        ("fig-backend-categories.png", "Figure 4.10: Backend API – Crime Categories (GET /api/categories)"),
        ("fig-backend-landing-stats.png", "Figure 4.11: Backend API – Landing Statistics (GET /api/landing/stats)"),
        ("fig-backend-track.png", "Figure 4.12: Backend API – Track Case JSON (GET /api/reports/track/{caseId})"),
        ("fig-backend-dashboard.png", "Figure 4.13: Backend API – Dashboard Data (GET /api/dashboard, authenticated)"),
    ]),
    ("DATABASE (MySQL / phpMyAdmin) SCREENSHOTS", [
        ("fig-db-structure.png", "Figure 4.14: MySQL Database ecrs_koforidua – Table Structure"),
        ("fig-db-reports.png", "Figure 4.15: MySQL – Reports Table Data"),
        ("fig-db-users.png", "Figure 4.16: MySQL – Users (Officers) Table"),
        ("fig-db-timeline.png", "Figure 4.17: MySQL – Case Timeline Table"),
    ]),
]


def add_figure(doc, image_path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.bold = True
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(11)
    doc.add_paragraph()


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    doc = Document(str(SRC))
    doc.add_page_break()

    title = doc.add_heading("APPENDIX: SYSTEM SCREENSHOTS", level=1)
    intro = doc.add_paragraph(
        "This appendix presents screenshots of the frontend interfaces, Laravel REST API "
        "backend responses, and MySQL database tables for the Electronic Crime Reporting System (ECRS)."
    )
    for run in intro.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    missing = []
    for section_title, figures in SECTIONS:
        doc.add_heading(section_title, level=2)
        for filename, caption in figures:
            path = SHOTS / filename
            if path.exists():
                add_figure(doc, path, caption)
            else:
                missing.append(filename)
                note = doc.add_paragraph(f"[Screenshot pending: {filename}]")
                for run in note.runs:
                    run.font.name = "Times New Roman"
                    run.italic = True

    doc.save(OUT)
    doc.save(DL)
    print(f"Saved: {OUT}")
    print(f"Saved: {DL}")
    if missing:
        print("Missing:", ", ".join(missing))


if __name__ == "__main__":
    main()
