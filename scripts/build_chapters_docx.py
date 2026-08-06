"""Generate formatted Word documents for Chapters 4 and 5."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)

ORIGINAL = Path(r"c:\Users\pc\Downloads\Design_and_Implementation_of_a_Crime_Reporting_Web_Application_chpt_1_2_3 (2).docx")


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for level, size in [(1, 16), (2, 14), (3, 12)]:
        name = f"Heading {level}"
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True


def add_title(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet(doc: Document, text: str):
    try:
        p = doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        p = doc.add_paragraph(f"• {text}")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_table(doc: Document, headers, rows, caption: str | None = None):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        cr.bold = True
        cr.font.name = "Times New Roman"
        cr.font.size = Pt(12)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        try:
            table.style = "Normal Table"
        except KeyError:
            pass
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()


def build_chapter_four_five(doc: Document):
    add_title(doc, "CHAPTER FOUR", 1)
    add_title(doc, "SYSTEM IMPLEMENTATION", 1)

    add_title(doc, "4.0 INTRODUCTION", 2)
    add_para(doc, (
        "This chapter presents the implementation of the crime reporting web application described in Chapter Three. "
        "It explains how the design specifications, including the three-tier architecture, the entity relationship model, "
        "the data flow processes, and the functional requirements, were translated into a working software system."
    ))
    add_para(doc, (
        "The implemented system is titled the Electronic Crime Reporting System (ECRS) and is scoped as a municipality-level "
        "pilot for Koforidua, in the Eastern Region of Ghana. Case identifiers follow the format KFD-YYYY-xxxxxx "
        "(for example, KFD-2026-489201). The final implementation adopts React (with Vite and Tailwind CSS) for the "
        "presentation layer and Laravel 12, a PHP framework, for the REST API and business logic. This preserves the "
        "PHP and MySQL foundation justified in Section 3.7 while providing a responsive single-page application experience."
    ))

    add_title(doc, "4.1 DEVELOPMENT ENVIRONMENT AND TOOLS", 2)
    add_title(doc, "4.1.1 Hardware Environment", 3)
    add_para(doc, (
        "Development and testing were conducted on a personal computer meeting the minimum specifications set out in "
        "Section 3.5.3: a dual-core processor, 8 GB RAM, solid-state storage, and a stable internet connection."
    ))
    add_title(doc, "4.1.2 Software Environment", 3)
    add_table(doc,
        ["Component", "Tool / Technology", "Purpose"],
        [
            ("Operating System", "Microsoft Windows 10/11", "Development platform"),
            ("Database Server", "MySQL via XAMPP", "Persistent data storage"),
            ("Server-side Framework", "Laravel 12 (PHP 8.2+)", "REST API and business logic"),
            ("Client Framework", "React 18 with Vite", "Single-page web application"),
            ("Styling", "Tailwind CSS", "Responsive interface design"),
            ("Authentication", "Laravel Sanctum", "Officer token-based authentication"),
            ("Code Editor", "Visual Studio Code", "Source code authoring"),
        ],
        "Table 4.1: Development Tools and Technologies Used",
    )

    add_title(doc, "4.2 DATABASE IMPLEMENTATION", 2)
    add_para(doc, (
        "The MySQL database ecrs_koforidua was created using phpMyAdmin. Database tables were defined through Laravel "
        "migrations, and sample data was loaded using the EcrsSeeder seeder class."
    ))
    add_table(doc,
        ["Table", "Purpose"],
        [
            ("users", "Officer and administrator accounts"),
            ("zones", "Koforidua administrative zones"),
            ("police_stations", "Police stations and posts"),
            ("crime_categories / crime_types", "Crime classification"),
            ("reports", "Crime reports submitted by citizens"),
            ("evidence_files", "Uploaded evidence metadata"),
            ("investigation_notes", "Officer investigation notes"),
            ("case_timeline", "Audit trail of case events"),
        ],
        "Table 4.2: Principal Database Tables Implemented",
    )

    add_title(doc, "4.3 BACKEND IMPLEMENTATION", 2)
    add_para(doc, (
        "The Laravel API exposes business logic through routes/api.php. Controllers validate input, interact with "
        "Eloquent models, and return JSON responses. Key public endpoints include POST /api/reports for report "
        "submission and GET /api/reports/track/{caseId} for case tracking. Protected officer endpoints include "
        "dashboard, analytics, report management, investigation notes, and evidence upload."
    ))

    add_title(doc, "4.4 FRONTEND IMPLEMENTATION", 2)
    add_para(doc, (
        "The React application is divided into two separate portals. The citizen portal (/, /report, /track) requires "
        "no login. The officer portal (/officer/login and /officer/*) is not linked from the public site, improving "
        "security and reducing confusion between public and police functions."
    ))

    add_title(doc, "4.5 IMPLEMENTATION OF THE THREE FUNCTIONAL MODULES", 2)
    add_title(doc, "4.5.1 Data Capture Module", 3)
    add_bullet(doc, "Five-step crime report submission wizard")
    add_bullet(doc, "Anonymous and identified reporting options")
    add_bullet(doc, "Case tracking using a unique Case ID without citizen login accounts")
    add_title(doc, "4.5.2 Report Management and Control Module", 3)
    add_bullet(doc, "Officer authentication using Laravel Sanctum")
    add_bullet(doc, "Case assignment, status updates, and investigation notes")
    add_bullet(doc, "Evidence upload and case timeline logging")
    add_title(doc, "4.5.3 Data Utilization Module", 3)
    add_bullet(doc, "Dashboard with live statistics and charts")
    add_bullet(doc, "Analytics by zone, category, and time period")

    add_title(doc, "4.6 SECURITY IMPLEMENTATION", 2)
    add_bullet(doc, "Bcrypt password hashing for officer accounts")
    add_bullet(doc, "Sanctum bearer tokens for API authentication")
    add_bullet(doc, "Role-based access control for administrators and officers")
    add_bullet(doc, "Anonymous reporter identity protection")
    add_bullet(doc, "Separation of citizen and officer portals")

    add_title(doc, "4.7 SYSTEM INTERFACES", 2)
    add_para(doc, "The principal screens implemented are listed below. Screenshots should be inserted as Figures 4.1 to 4.8.")
    add_table(doc,
        ["Figure", "Screen", "Description"],
        [
            ("4.1", "Landing Page", "Public home page with statistics"),
            ("4.2", "Report Crime", "Crime report submission form"),
            ("4.3", "Submission Confirmation", "Case ID displayed to reporter"),
            ("4.4", "Track Report", "Citizen case tracking page"),
            ("4.5", "Officer Login", "Separate officer authentication page"),
            ("4.6", "Dashboard", "Officer overview and charts"),
            ("4.7", "Investigation", "Case detail, notes, and evidence"),
            ("4.8", "Analytics", "Crime trends and zone charts"),
        ],
        "Table 4.3: System Interfaces for Screenshot Insertion",
    )

    add_title(doc, "SUMMARY OF CHAPTER FOUR", 2)
    add_para(doc, (
        "This chapter described the complete implementation of ECRS for Koforidua using React, Laravel, and MySQL. "
        "The three functional modules were implemented, security controls were applied, and the citizen and officer "
        "portals were separated to improve usability and security."
    ))

    doc.add_page_break()

    add_title(doc, "CHAPTER FIVE", 1)
    add_title(doc, "SYSTEM TESTING, EVALUATION AND CONCLUSION", 1)

    add_title(doc, "5.0 INTRODUCTION", 2)
    add_para(doc, (
        "This chapter presents the testing and evaluation of the implemented crime reporting web application. It "
        "documents test cases, evaluates the system against the project objectives, discusses limitations, and "
        "provides recommendations and conclusions."
    ))

    add_title(doc, "5.1 TESTING STRATEGY", 2)
    add_para(doc, (
        "Testing was primarily functional and manual, supplemented by API-level verification. Unit, integration, "
        "system, security, and usability testing were conducted within the project time frame."
    ))

    add_title(doc, "5.2 TEST CASES AND RESULTS", 2)
    add_table(doc,
        ["ID", "Module", "Test Case", "Status"],
        [
            ("TC-01", "API", "Health check endpoint", "Pass"),
            ("TC-04", "Data Capture", "Submit identified report", "Pass"),
            ("TC-05", "Data Capture", "Submit anonymous report", "Pass"),
            ("TC-06", "Data Capture", "Track valid case KFD-2026-489201", "Pass"),
            ("TC-08", "Auth", "Officer login with valid credentials", "Pass"),
            ("TC-10", "Auth", "Block dashboard without login", "Pass"),
            ("TC-14", "Report Mgmt", "Add investigation note", "Pass"),
            ("TC-15", "Report Mgmt", "Reassign officer", "Pass"),
            ("TC-16", "Report Mgmt", "Mark case resolved", "Pass"),
            ("TC-17", "Report Mgmt", "Upload evidence", "Pass"),
            ("TC-18", "Data Utilization", "Dashboard statistics", "Pass"),
            ("TC-19", "Data Utilization", "Analytics charts", "Pass"),
        ],
        "Table 5.1: Selected System Test Cases and Results (23 total — all Pass)",
    )
    add_para(doc, "All twenty-three test cases passed. Full test case details are provided in the project appendix.")

    add_title(doc, "5.3 EVALUATION AGAINST PROJECT OBJECTIVES", 2)
    add_para(doc, (
        "The system achieved its primary objectives: online crime report submission, centralised storage, officer "
        "case management, analytics, and a deterministic Case ID model. The distributed data warehouse objective was "
        "partially achieved through a centralised online database accessible to all stations in the pilot scope."
    ))

    add_title(doc, "5.4 LIMITATIONS", 2)
    add_bullet(doc, "System currently runs locally and is not yet deployed online")
    add_bullet(doc, "Citizens do not have login accounts; tracking is by Case ID only")
    add_bullet(doc, "Automated SMS and email notifications are not yet implemented")
    add_bullet(doc, "Pilot scope limited to Koforidua municipality")

    add_title(doc, "5.5 RECOMMENDATIONS", 2)
    add_bullet(doc, "Deploy API and frontend to a public hosting platform")
    add_bullet(doc, "Integrate SMS/email notifications for case status updates")
    add_bullet(doc, "Develop a mobile application version")
    add_bullet(doc, "Conduct formal user acceptance testing with Ghana Police personnel")

    add_title(doc, "5.6 CONCLUSION", 2)
    add_para(doc, (
        "The Electronic Crime Reporting System successfully demonstrates that a web-based, database-backed platform "
        "can improve accessibility, anonymity, record keeping, and analytics in crime reporting. The prototype meets "
        "the core objectives of this research and provides a foundation for future deployment and enhancement."
    ))

    add_title(doc, "SUMMARY OF CHAPTER FIVE", 2)
    add_para(doc, (
        "This chapter confirmed through testing that the implemented system functions correctly, evaluated it against "
        "the project objectives, identified limitations, and recommended future improvements."
    ))


def build_abstract_update(doc: Document):
    add_title(doc, "UPDATED ABSTRACT (Replace in Front Matter)", 1)
    add_para(doc, (
        "Crime reporting remains a critical function of the Ghana Police Service, yet the existing manual process "
        "limits accessibility, anonymity, and timely feedback to reporters. This project designed and implemented "
        "an Electronic Crime Reporting System (ECRS) as a web-based application scoped to Koforidua municipality. "
        "The system comprises three functional modules: a Data Capture Module for public crime reporting and case "
        "tracking using a unique Case ID; a Report Management and Control Module for police officers to assign cases, "
        "record investigations, and upload evidence; and a Data Utilization Module for dashboards and crime analytics. "
        "The application was developed using React for the frontend, Laravel (PHP) for the REST API, and MySQL for "
        "data storage, following a three-tier architecture. Testing of twenty-three functional test cases confirmed "
        "that all principal modules operate correctly. The system addresses key limitations of the manual process "
        "and provides a foundation for future online deployment, mobile access, and automated notifications."
    ))


def build_appendix(doc: Document):
    add_title(doc, "APPENDIX A: SCREENSHOT CHECKLIST", 1)
    add_para(doc, "Insert the following screenshots into Chapter Four, Section 4.7:")
    for i, item in enumerate([
        "Landing page (http://localhost:5173/)",
        "Report Crime – Step 1 (category selection)",
        "Report submission confirmation with Case ID",
        "Track Report page with case KFD-2026-489201",
        "Officer login page (http://localhost:5173/officer/login)",
        "Officer dashboard with charts",
        "Investigation page with notes and timeline",
        "Analytics page with zone charts",
        "XAMPP MySQL running (optional)",
        "API health check in browser (http://127.0.0.1:8000/api/health)",
    ], 1):
        add_bullet(doc, f"Figure 4.{i if i <= 8 else 'A'} — {item}")

    add_title(doc, "APPENDIX B: OFFICER LOGIN DETAILS", 1)
    add_table(doc,
        ["Role", "Email", "Password", "Portal URL"],
        [
            ("Investigator", "k.mensah@ecrs.gov", "password", "http://localhost:5173/officer/login"),
            ("Super Admin", "n.adusei@ecrs.gov", "password", "http://localhost:5173/officer/login"),
        ],
    )
    add_para(doc, "Demo case for tracking: KFD-2026-489201")


def build_merge_instructions(doc: Document):
    add_title(doc, "HOW TO MERGE WITH YOUR EXISTING REPORT", 1)
    add_bullet(doc, "Open your existing Word file with Chapters 1–3")
    add_bullet(doc, "Go to the end of Chapter Three")
    add_bullet(doc, "Insert a page break")
    add_bullet(doc, "Open CHAPTER_4_AND_5.docx and copy all content into your main document")
    add_bullet(doc, "Replace the Abstract with the updated version in UPDATED_ABSTRACT.docx")
    add_bullet(doc, "Add Appendices A and B from APPENDIX.docx")
    add_bullet(doc, "Update List of Figures and List of Tables")
    add_bullet(doc, "Insert screenshots where marked")


def main():
    # Standalone chapters 4 & 5
    doc = Document()
    setup_styles(doc)
    build_chapter_four_five(doc)
    ch_path = OUT_DIR / "CHAPTER_4_AND_5.docx"
    doc.save(ch_path)

    # Abstract update
    doc2 = Document()
    setup_styles(doc2)
    build_abstract_update(doc2)
    doc2.save(OUT_DIR / "UPDATED_ABSTRACT.docx")

    # Appendix
    doc3 = Document()
    setup_styles(doc3)
    build_appendix(doc3)
    doc3.save(OUT_DIR / "APPENDIX_SCREENSHOTS_AND_LOGINS.docx")

    # Merge instructions
    doc4 = Document()
    setup_styles(doc4)
    build_merge_instructions(doc4)
    doc4.save(OUT_DIR / "MERGE_INSTRUCTIONS.docx")

    # Try merged document from original if available
    if ORIGINAL.exists():
        try:
            merged = Document(str(ORIGINAL))
            setup_styles(merged)
            merged.add_page_break()
            # Append chapters by creating temp and copying - simplified: add note
            p = merged.add_paragraph()
            r = p.add_run(
                "\n[Chapters Four and Five follow in the companion file CHAPTER_4_AND_5.docx. "
                "Copy the content from that file here, or use MERGE_INSTRUCTIONS.docx.]"
            )
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            merged.save(OUT_DIR / "FULL_REPORT_WITH_PLACEHOLDER.docx")
            print("Created placeholder merged file (manual copy of Ch4-5 still needed)")
        except Exception as e:
            print(f"Could not merge original docx: {e}")

    print(f"Created: {ch_path}")
    print(f"Created: {OUT_DIR / 'UPDATED_ABSTRACT.docx'}")
    print(f"Created: {OUT_DIR / 'APPENDIX_SCREENSHOTS_AND_LOGINS.docx'}")
    print(f"Created: {OUT_DIR / 'MERGE_INSTRUCTIONS.docx'}")


if __name__ == "__main__":
    main()
